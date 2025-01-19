"""
Module for building Word documents (.docx) using student and common configurations.

This module processes template Word files, fills in the required fields, and generates formatted
documents for students based on their disciplines, diploma themes, and other configurations.

Functions:
- build_docx: Generates a single Word document for a specific student.
- build_statements: Builds multiple Word documents for a list of students.

Dependencies:
- python-docx: For creating and modifying Word documents.
- os: For handling file paths.
- backend.classes.student_config: Student configuration class.
- backend.classes.common_config: Common configuration class.
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from backend.classes.student_config import StudentConfig
from backend.classes.common_config import CommonConfig
import os


def build_docx(
        template_path: str,
        save_directory_path: str,
        student_config: StudentConfig,
        common_config: CommonConfig,
) -> None:
    """
    Generates a single Word document (.docx) for a student based on the provided template.

    :param template_path: The file path to the Word document template.
    :type template_path: str
    :param save_directory_path: The directory path where the generated document will be saved.
    :type save_directory_path: str
    :param student_config: Configuration containing student's disciplines, full name, and diploma theme.
    :type student_config: StudentConfig
    :param common_config: Configuration containing common information such as dates and specialty details.
    :type common_config: CommonConfig
    :return: None
    """
    document = Document(template_path)

    paragraph_1 = document.paragraphs[5]
    paragraph_1.clear()
    run_1_1 = paragraph_1.add_run(student_config.full_name)
    run_1_1.font.size = Pt(13.5)
    run_1_1.underline = True

    paragraph_2 = document.paragraphs[10]
    paragraph_2.clear()
    run_2_1 = paragraph_2.add_run(f"с {common_config.start_date_day} {common_config.start_date_month} 20")
    run_2_1.font.size = Pt(12)
    run_2_2 = paragraph_2.add_run(f"{common_config.start_date_year}")
    run_2_2.font.size = Pt(12)
    run_2_2.underline = True
    run_2_3 = paragraph_2.add_run(f" г. по {common_config.end_date_day} {common_config.end_date_month} 20")
    run_2_3.font.size = Pt(12)
    run_2_4 = paragraph_2.add_run(f"{common_config.end_date_year}")
    run_2_4.font.size = Pt(12)
    run_2_4.underline = True
    run_2_5 = paragraph_2.add_run(f" г.")
    run_2_5.font.size = Pt(12)

    paragraph_3 = document.paragraphs[11]
    paragraph_3.clear()
    run_3_1 = paragraph_3.add_run(f"по специальности ")
    run_3_1.font.size = Pt(12)
    run_3_2 = paragraph_3.add_run(f"{common_config.speciality_code} «{common_config.speciality_name}»")
    run_3_2.font.size = Pt(12)
    run_3_2.underline = True

    paragraph_4 = document.paragraphs[12]
    paragraph_4.clear()
    run_4_1 = paragraph_4.add_run(f"направлению специальности ")
    run_4_1.font.size = Pt(12)
    run_4_2 = paragraph_4.add_run(f"{common_config.speciality_area_code} «{common_config.speciality_area_name}»")
    run_4_2.font.size = Pt(12)
    run_4_2.underline = True

    paragraph_5 = document.paragraphs[19]
    paragraph_5.clear()
    run_5_1 = paragraph_5.add_run(f"Выполнил(а) дипломный проект на тему: ")
    run_5_1.font.size = Pt(11)
    run_5_2 = paragraph_5.add_run(f"«{student_config.diploma_theme}»")
    run_5_2.font.size = Pt(11)
    run_5_2.underline = True

    paragraph_6 = document.paragraphs[31]
    paragraph_6.clear()
    run_6_1 = paragraph_6.add_run(f"г. Могилев «{common_config.statement_date_day}»"
                                  f" {common_config.statement_date_month}"
                                  f" {common_config.statement_date_year}\tРегистрационный № ____")
    run_6_1.font.size = Pt(11)

    table1 = document.tables[0]
    for discipline in student_config.regular_disciplines:
        new_row = table1.add_row()

        cell_1 = new_row.cells[0]
        cell_paragraph_1 = cell_1.paragraphs[0]
        cell_run_1 = cell_paragraph_1.add_run(discipline[0])
        cell_run_1.font.size = Pt(10)

        cell_2 = new_row.cells[1]
        cell_paragraph_2 = cell_2.paragraphs[0]
        cell_run_2 = cell_paragraph_2.add_run(discipline[1])
        cell_run_2.font.size = Pt(10)

        cell_3 = new_row.cells[2]
        cell_paragraph_3 = cell_3.paragraphs[0]
        cell_run_3 = cell_paragraph_3.add_run(discipline[2])
        cell_run_3.font.size = Pt(10)

        cell_paragraph_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell_paragraph_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table2 = document.tables[1]
    for discipline in student_config.course_work_disciplines:
        new_row = table2.add_row()

        cell_1 = new_row.cells[0]
        cell_paragraph_1 = cell_1.paragraphs[0]
        cell_run_1 = cell_paragraph_1.add_run(discipline[0])
        cell_run_1.font.size = Pt(10)

        cell_2 = new_row.cells[1]
        cell_paragraph_2 = cell_2.paragraphs[0]
        cell_run_2 = cell_paragraph_2.add_run(discipline[1])
        cell_run_2.font.size = Pt(10)

        cell_3 = new_row.cells[2]
        cell_paragraph_3 = cell_3.paragraphs[0]
        cell_run_3 = cell_paragraph_3.add_run(discipline[2])
        cell_run_3.font.size = Pt(10)

        cell_paragraph_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell_paragraph_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for discipline in student_config.course_project_disciplines:
        new_row = table2.add_row()

        cell_1 = new_row.cells[0]
        cell_paragraph_1 = cell_1.paragraphs[0]
        cell_run_1 = cell_paragraph_1.add_run(discipline[0])
        cell_run_1.font.size = Pt(10)

        cell_2 = new_row.cells[1]
        cell_paragraph_2 = cell_2.paragraphs[0]
        cell_run_2 = cell_paragraph_2.add_run(discipline[1])
        cell_run_2.font.size = Pt(10)

        cell_3 = new_row.cells[2]
        cell_paragraph_3 = cell_3.paragraphs[0]
        cell_run_3 = cell_paragraph_3.add_run(discipline[2])
        cell_run_3.font.size = Pt(10)

        cell_paragraph_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell_paragraph_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table3 = document.tables[2]
    for discipline in student_config.practice_disciplines:
        new_row = table3.add_row()

        cell_1 = new_row.cells[0]
        cell_paragraph_1 = cell_1.paragraphs[0]
        cell_run_1 = cell_paragraph_1.add_run(discipline[0])
        cell_run_1.font.size = Pt(10)

        cell_2 = new_row.cells[1]
        cell_paragraph_2 = cell_2.paragraphs[0]
        cell_run_2 = cell_paragraph_2.add_run(discipline[1])
        cell_run_2.font.size = Pt(10)

        cell_3 = new_row.cells[2]
        cell_paragraph_3 = cell_3.paragraphs[0]
        cell_run_3 = cell_paragraph_3.add_run(discipline[2])
        cell_run_3.font.size = Pt(10)

        cell_paragraph_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell_paragraph_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    file_name = f"Выписка_{student_config.full_name}.docx"
    save_path = os.path.join(save_directory_path, file_name)
    document.save(save_path)


def build_statements(
        template_path: str,
        save_directory_path: str,
        students_configs: list[StudentConfig],
        common_config: CommonConfig,
) -> None:
    """
    Generates multiple Word documents (.docx) for a list of students using a provided template.

    :param template_path: The file path to the Word document template.
    :type template_path: str
    :param save_directory_path: The directory path where all generated documents will be saved.
    :type save_directory_path: str
    :param students_configs: A list of StudentConfig objects containing individual student data.
    :type students_configs: list[StudentConfig]
    :param common_config: Common configuration containing shared details like dates and specialties.
    :type common_config: CommonConfig
    :return: None
    """
    for student_config in students_configs:
        build_docx(template_path,
                   save_directory_path,
                   student_config,
                   common_config)
